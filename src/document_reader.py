"""
Document Reader Utility

Håndterer læsning af forskellige dokumentformater (TXT, PDF, Word).
"""

from pathlib import Path
from typing import Tuple
import pypdf
from docx import Document


class DocumentReader:
    """Læser forskellige dokumentformater og returnerer tekst."""

    @staticmethod
    def read_file(file_path: Path) -> Tuple[str, str]:
        """
        Læser en fil og returnerer indholdet som tekst.

        Args:
            file_path: Sti til filen der skal læses

        Returns:
            Tuple af (tekst_indhold, filformat)

        Raises:
            ValueError: Hvis filformatet ikke understøttes
            Exception: Ved læsefejl
        """
        file_extension = file_path.suffix.lower()

        if file_extension == '.txt':
            return DocumentReader._read_txt(file_path), 'txt'
        elif file_extension == '.pdf':
            return DocumentReader._read_pdf(file_path), 'pdf'
        elif file_extension in ['.docx', '.doc']:
            return DocumentReader._read_docx(file_path), 'docx'
        else:
            raise ValueError(f"Ikke-understøttet filformat: {file_extension}")

    @staticmethod
    def _read_txt(file_path: Path) -> str:
        """Læser en TXT fil."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Fallback til latin-1 hvis UTF-8 fejler
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()

    @staticmethod
    def _read_pdf(file_path: Path) -> str:
        """Læser en PDF fil."""
        text = []

        try:
            with open(file_path, 'rb') as f:
                pdf_reader = pypdf.PdfReader(f)

                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text.append(page.extract_text())

            return '\n\n'.join(text)
        except Exception as e:
            raise Exception(f"Fejl ved læsning af PDF: {e}")

    @staticmethod
    def _read_docx(file_path: Path) -> str:
        """Læser en Word dokument (.docx)."""
        try:
            doc = Document(file_path)
            text = []

            # Læs alle paragraffer
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)

            # Læs også tabeller hvis der er nogen
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text.append(' | '.join(row_text))

            return '\n\n'.join(text)
        except Exception as e:
            raise Exception(f"Fejl ved læsning af Word dokument: {e}")

    @staticmethod
    def write_file(file_path: Path, content: str, file_format: str = 'txt') -> None:
        """
        Skriver indhold til en fil i det angivne format.

        Args:
            file_path: Sti til filen der skal skrives
            content: Tekst indhold der skal skrives
            file_format: Format ('txt', 'pdf', 'docx')
        """
        if file_format == 'txt':
            DocumentReader._write_txt(file_path, content)
        elif file_format == 'docx':
            DocumentReader._write_docx(file_path, content)
        elif file_format == 'pdf':
            # PDF skrivning er mere kompleks - for nu gemmer vi som TXT
            # Vi kan tilføje PDF generation senere hvis nødvendigt
            txt_path = file_path.with_suffix('.txt')
            DocumentReader._write_txt(txt_path, content)
        else:
            DocumentReader._write_txt(file_path, content)

    @staticmethod
    def _write_txt(file_path: Path, content: str) -> None:
        """Skriver til TXT fil."""
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)

    @staticmethod
    def _write_docx(file_path: Path, content: str) -> None:
        """Skriver til Word dokument."""
        doc = Document()

        # Split content into paragraphs
        paragraphs = content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())

        doc.save(file_path)

    @staticmethod
    def get_supported_extensions() -> list:
        """Returnerer liste af understøttede fil-endelser."""
        return ['.txt', '.pdf', '.docx', '.doc']

    @staticmethod
    def is_supported(file_path: Path) -> bool:
        """Tjek om en fil er i et understøttet format."""
        return file_path.suffix.lower() in DocumentReader.get_supported_extensions()
